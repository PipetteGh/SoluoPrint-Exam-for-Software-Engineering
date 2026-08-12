import os
import zlib
import base64
import urllib.request
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_kroki_url(diagram_type, diagram_text):
    data = base64.urlsafe_b64encode(zlib.compress(diagram_text.encode('utf-8'), 9)).decode('ascii')
    return f"https://kroki.io/{diagram_type}/png/{data}"

def download_image(url, filename):
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open(filename, 'wb') as out_file:
        out_file.write(response.read())

def set_font(run, name='Times New Roman', size=12, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_font(run, size=16 if level == 1 else 14, bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    heading.paragraph_format.space_after = Pt(12)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(12)
    return p

def main():
    doc = Document()
    
    # Update default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Software Requirements Specification (SRS)\nSoluoPrint")
    set_font(title_run, size=24, bold=True)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(24)

    add_heading(doc, "1. Introduction", level=1)
    add_paragraph(doc, "The SoluoPrint system is a comprehensive print shop management application designed to handle customers, print jobs, financial transactions, and automated notifications (SMS). This document formalizes the requirements and architecture of the existing system as part of the CSCD 602 Advanced Software Engineering Project.")

    add_heading(doc, "2. Functional Requirements", level=1)
    reqs = [
        "FR1: The system shall allow users to register and create a company profile.",
        "FR2: The system shall enforce email verification for new accounts prior to granting access.",
        "FR3: The system shall allow authorized users to create, update, and manage print jobs with statuses (Pending, In Progress, Completed, Delivered).",
        "FR4: The system shall allow users to manage customer records and track customer balances.",
        "FR5: The system shall allow users to record payments against specific jobs or general customer balances.",
        "FR6: The system shall automatically send an SMS notification using SMSGH SMSAPI when a payment is received, formatted exactly as: 'Hello [Name], your payment of: [Currency] [Amount] has been received. Your remaining balance is [Currency] [Balance]. Regards: [Company Name] - [Phone] Powered by: Soluotech'.",
        "FR7: The system shall provide Role-Based Access Control (RBAC) supporting 'owner' and 'admin' roles.",
        "FR8: The system shall display analytics on a Dashboard (Total Revenue, Jobs, Outstanding Balances)."
    ]
    for r in reqs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(r)
        set_font(p.runs[0])
        p.paragraph_format.line_spacing = 1.5

    add_heading(doc, "3. Non-Functional Requirements", level=1)
    nfrs = [
        "NFR1 (Security): Passwords and authentication must be handled securely via Supabase Auth. Passwords shall not be stored in plain text.",
        "NFR2 (Performance): The dashboard shall load within 2 seconds under normal conditions.",
        "NFR3 (Usability): The interface must be fully responsive, scaling appropriately for desktop and mobile browsers.",
        "NFR4 (Reliability): The system must rely on Row-Level Security (RLS) to strictly isolate company data in a multi-tenant environment."
    ]
    for r in nfrs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(r)
        set_font(p.runs[0])
        p.paragraph_format.line_spacing = 1.5

    # Diagrams
    add_heading(doc, "4. System Architecture & Diagrams", level=1)
    
    # 4.1 System Context
    add_heading(doc, "4.1 System Context Diagram", level=2)
    add_paragraph(doc, "The context diagram illustrates how the SoluoPrint system interacts with external entities (Users, Supabase Backend, and SMSGH SMSAPI).")
    
    context_uml = """
    @startuml
    !theme plain
    skinparam componentStyle rectangle
    actor "Print Shop Owner/Admin" as Admin
    actor "Customer" as Cust
    
    node "SoluoPrint React Application" as App
    
    cloud "Supabase Backend" {
      [PostgreSQL Database]
      [Auth Service]
    }
    
    cloud "SMSGH SMSAPI" as SMS
    
    Admin --> App : Manages Jobs & Payments
    App --> [Auth Service] : Authenticates
    App --> [PostgreSQL Database] : Reads/Writes Data
    App --> SMS : Triggers Notifications
    SMS --> Cust : Delivers SMS
    @enduml
    """
    
    context_url = create_kroki_url("plantuml", context_uml)
    download_image(context_url, "context.png")
    doc.add_picture("context.png", width=Inches(5))
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 4.2 Use Case Diagram
    add_heading(doc, "4.2 Use Case Diagram", level=2)
    add_paragraph(doc, "The following diagram highlights the primary use cases available to an authenticated user.")
    
    usecase_uml = """
    @startuml
    !theme plain
    left to right direction
    actor "Shop User" as User
    
    rectangle SoluoPrint {
      usecase "Manage Customers" as UC1
      usecase "Create Print Job" as UC2
      usecase "Record Payment" as UC3
      usecase "Configure SMS Settings" as UC4
      usecase "View Analytics Dashboard" as UC5
    }
    
    User --> UC1
    User --> UC2
    User --> UC3
    User --> UC4
    User --> UC5
    @enduml
    """
    
    uc_url = create_kroki_url("plantuml", usecase_uml)
    download_image(uc_url, "usecase.png")
    doc.add_picture("usecase.png", width=Inches(5))
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 4.3 ER Diagram
    add_heading(doc, "4.3 Entity Relationship Diagram", level=2)
    add_paragraph(doc, "The core data model for the application, mapping relationships between companies, profiles, customers, jobs, and payments.")
    
    er_uml = """
    @startuml
    !theme plain
    entity "companies" as C {
      * id : uuid
      --
      name : text
      currency : text
    }
    entity "profiles" as P {
      * id : uuid
      --
      company_id : uuid
      role : text
    }
    entity "customers" as Cust {
      * id : uuid
      --
      company_id : uuid
      balance : numeric
    }
    entity "print_jobs" as J {
      * id : uuid
      --
      customer_id : uuid
      total_price : numeric
      status : text
    }
    entity "payments" as Pay {
      * id : uuid
      --
      customer_id : uuid
      job_id : uuid
      amount : numeric
    }
    
    C ||--o{ P : has
    C ||--o{ Cust : manages
    Cust ||--o{ J : places
    J ||--o{ Pay : receives
    Cust ||--o{ Pay : receives
    @enduml
    """
    
    er_url = create_kroki_url("plantuml", er_uml)
    download_image(er_url, "er.png")
    doc.add_picture("er.png", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save("SRS.docx")
    print("SRS.docx generated successfully!")
    
    # Cleanup images
    os.remove("context.png")
    os.remove("usecase.png")
    os.remove("er.png")

if __name__ == "__main__":
    main()
